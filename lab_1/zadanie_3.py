from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import scipy.fftpack as fft
import sounddevice as sd
import soundfile as sf


def plotAudio(Signal, fs, fsize, axs, TimeMargin=[0, 0.02]):

    x_time = np.arange(len(Signal)) / fs
    x_frequency = np.arange(0, fs / 2, fs / fsize)
    spectrum_halved = fft.fft(Signal, fsize)[: fsize // 2]
    spectrum_dB = 20 * np.log10(np.abs(spectrum_halved))

    axs[0].plot(x_time, Signal)
    axs[0].set_title("Audio Signal")
    axs[0].set_xlabel("Time (s)")
    axs[0].set_ylabel("Amplitude")
    axs[0].set_xlim(TimeMargin)
    axs[0].grid()

    axs[1].plot(x_frequency, spectrum_dB)
    axs[1].set_title("Spectrum")
    axs[1].set_xlabel("Frequency (Hz)")
    axs[1].set_ylabel("Amplitude")
    axs[1].grid()

    max_amplitude_index = np.argmax(spectrum_dB)
    peak_amplitude = spectrum_dB[max_amplitude_index]
    peak_frequency = x_frequency[max_amplitude_index]

    return peak_frequency, peak_amplitude


def generate_report(files, fsizes, output_file="report.docx"):
    document = Document()
    document.add_heading("Analiza sinusoidalnych sygnałów", 0)

    for file in files:
        document.add_heading(f"Plik {file}", level=2)

        # ----------------------------------------------------------------
        signal, fs = sf.read(file)
        # ----------------------------------------------------------------

        for fsize in fsizes:
            document.add_heading(f"Rozmiar okna FFT: {fsize}", level=3)

            fig, axs = plt.subplots(2, 1, figsize=(10, 7))

            # ----------------------------------------------------------------
            peak_frequency, peak_amplitude = plotAudio(signal, fs, fsize, axs)
            # ----------------------------------------------------------------

            fig.suptitle(f"Rozmiar okna FFT: {fsize}")
            fig.tight_layout(pad=1.5)
            memfile = BytesIO()
            fig.savefig(memfile)

            document.add_picture(memfile, width=Inches(6))

            memfile.close()

            document.add_paragraph(
                f"Największa amplituda: {peak_amplitude:.3f} dB dla częstotliwości {peak_frequency:.3f} Hz"
            )

    document.save(output_file)
    print("Report generated")


def test(files, fsizes):
    for file in files:
        signal, fs = sf.read(file)

        for fsize in fsizes:
            fig, axs = plt.subplots(2, 1, figsize=(10, 7))
            plotAudio(signal, fs, fsize, axs)


if __name__ == "__main__":
    fsizes = [2**8, 2**12, 2**16]
    files = [
        "SOUND_SIN/sin_60Hz.wav",
        "SOUND_SIN/sin_440Hz.wav",
        "SOUND_SIN/sin_8000Hz.wav",
    ]

    # generate_report(["SOUND_SIN/sin_8000Hz.wav"], fsize)

    # fig, axs = plt.subplots(2, 1)
    # signal, fs = sf.read(files[1])
    # plotAudio(signal, fs, fsize[0], axs)
    # plt.show()

    # for file in files:
    #     signal, fs = sf.read(file)

    #     for fsize in fsizes:
    #         fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    #         plotAudio(signal, fs, fsize, axs)
    #         plt.show()

    generate_report(files, fsizes, "zadanie_3.docx")

    # fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    # test(files, fsizes)
