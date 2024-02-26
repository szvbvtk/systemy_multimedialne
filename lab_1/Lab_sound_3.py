from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import scipy.fftpack as fft
import sounddevice as sd
import soundfile as sf


def plotAudio(Signal, fs, fsize, axs, TimeMargin=[0, 0.02]):

    x = np.arange(len(Signal)) / fs
    axs[0].plot(x, Signal)
    axs[0].set_title("Audio Signal")
    axs[0].set_xlabel("Time (s)")
    axs[0].set_ylabel("Amplitude")
    axs[0].set_xlim(TimeMargin)
    axs[0].grid()

    spectrum = fft.fft(Signal, fsize)
    spectrum_dB_half = 20 * np.log10(np.abs(spectrum[: fsize // 2]))
    x_freqs = np.arange(0, fs / 2, fs / fsize)
    axs[1].plot(x_freqs, spectrum_dB_half)
    axs[1].set_title("Spectrum")
    axs[1].set_xlabel("Frequency (Hz)")
    axs[1].set_ylabel("Amplitude")
    axs[1].grid()

    max_amplitude_index = np.argmax(spectrum_dB_half)
    max_amplitude = spectrum_dB_half[max_amplitude_index]
    max_frequency = x_freqs[max_amplitude_index]

    # axs[1].annotate("Amplituda maks.: {:.2f} dB".format(max_amplitude), xy=(max_frequency, max_amplitude), xytext=(max_frequency, max_amplitude + 3))

    return max_frequency, max_amplitude
    return 0,0


def generate_report(files, fsizes):
    document = Document()
    document.add_heading("Analiza sinusoidalnych sygnałów", 0)

    for file in files:
        document.add_heading(f"Plik {file}", level=2)

        signal, fs = sf.read(file)

        for fsize in fsizes:
            document.add_heading(f"Rozmiar okna FFT: {fsize}", level=3)

            fig, axs = plt.subplots(2, 1, figsize=(10, 7))

            max_frequency, max_amplitude = plotAudio(signal, fs, fsize, axs)

            fig.suptitle(f"Rozmiar okna FFT: {fsize}")
            fig.tight_layout(pad=1.5)
            memfile = BytesIO()
            fig.savefig(memfile)

            document.add_picture(memfile, width=Inches(6))

            memfile.close()

            # Tu dodajesz dane tekstowe - wartości, wyjście funkcji etc.
            document.add_paragraph(f"Wartość losowa = {np.random.rand(1)}")

    document.save("report.docx")
    print("Report generated")

def test(files, fsizes):
    for file in files:
        signal, fs = sf.read(file)

        for fsize in fsizes:
            fig, axs = plt.subplots(2, 1, figsize=(10, 7))
            plotAudio(signal, fs, fsize, axs)
            

if __name__ == "__main__":
    fsizes = [2**8, 2**12, 2**16]
    files = [
        "SOUND_SIN/sin_60Hz.wav",
        "SOUND_SIN/sin_440Hz.wav",
        "SOUND_SIN/sin_8000Hz.wav",
    ]

    # generate_report(["SOUND_SIN/sin_8000Hz.wav"], fsize)

    # fig, axs = plt.subplots(2, 1)
    # signal, fs = sf.read(files[1])
    # plotAudio(signal, fs, fsize[0], axs)
    # plt.show()

    # for file in files:
    #     signal, fs = sf.read(file)

    #     for fsize in fsizes:
    #         fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    #         plotAudio(signal, fs, fsize, axs)
    #         plt.show()

    generate_report(files, fsizes)

    # fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    # test(files, fsizes)
