from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

document = Document()
document.add_heading('Zmień ten tytuł',0) # tworzenie nagłówków druga wartość to poziom nagłówka 


files=['sin60Hz.wav','sin440Hz.wav','sin8000Hz.wav']
Margins=[[0,0.02],[0.133,0.155]]
for file in files:
    document.add_heading('Plik - {}'.format(file),2)
    for i,Margin in enumerate(Margins):
        document.add_heading('Time margin {}'.format(Margin),3) # nagłówek sekcji, mozę być poziom wyżej
        fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
    
        ############################################################
        # Tu wykonujesz jakieś funkcje i rysujesz wykresy
        ############################################################
        
        fig.suptitle('Time margin {}'.format(Margin)) # Tytuł wykresu
        fig.tight_layout(pad=1.5) # poprawa czytelności 
        memfile = BytesIO() # tworzenie bufora
        fig.savefig(memfile) # z zapis do bufora 
        
    
        document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
        
        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1))) 
        ############################################################

document.save('report.docx') # zapis do pliku